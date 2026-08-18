from pathlib import Path

source_path = Path("GUIDE_UTILISATEUR_E_ADMINISTRATION.md")
out_path = Path("GUIDE_UTILISATEUR_E_ADMINISTRATION.docx")

try:
    from docx import Document
except Exception:
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

doc = Document()

for raw_line in source_path.read_text(encoding="utf-8").splitlines():
    line = raw_line.rstrip()

    if not line.strip():
        doc.add_paragraph("")
        continue

    if line.startswith("---"):
        doc.add_paragraph("=" * 60)
        continue

    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
        continue

    if line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=2)
        continue

    if line.startswith("### "):
        doc.add_heading(line[4:].strip(), level=3)
        continue

    if line.startswith("- "):
        doc.add_paragraph(line[2:].strip(), style="List Bullet")
        continue

    if line[:2].isdigit() and line[1:3] == ". ":
        doc.add_paragraph(line[3:].strip(), style="List Number")
        continue

    doc.add_paragraph(line)

doc.save(out_path)
print(f"Generated: {out_path}")
